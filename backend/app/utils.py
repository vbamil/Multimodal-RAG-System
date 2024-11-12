# File: backend/app/utils.py

import logging
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io
import re
from typing import List, Dict, Tuple
import spacy
import nltk
import unicodedata
import camelot
import pdfplumber
import tabula
import pandas as pd  # Required for Tabula-py
import docx  # For DOCX table extraction

from src.chunkers.text_chunker import chunk_text  # Ensure correct import path
from src.chunkers.batch_chunker import batch_chunk_text  # Ensure correct import path

# Initialize logger
logger = logging.getLogger(__name__)
logger.setLevel(logging.DEBUG)  # Set to DEBUG for detailed logs

# Create console handler with a higher log level
ch = logging.StreamHandler()
ch.setLevel(logging.DEBUG)  # Capture all logs at DEBUG level

# Create formatter and add it to the handlers
formatter = logging.Formatter('%(levelname)s:%(name)s:%(message)s')
ch.setFormatter(formatter)

# Add the handlers to the logger if not already added
if not logger.handlers:
    logger.addHandler(ch)

# Load spaCy model with error handling
try:
    # For entity extraction
    nlp = spacy.load("en_core_web_lg")  # Upgraded to a larger model for better NER accuracy
except OSError:
    logger.info("Downloading spaCy 'en_core_web_lg' model...")
    from spacy.cli import download
    download("en_core_web_lg")
    nlp = spacy.load("en_core_web_lg")

# Download NLTK punkt tokenizer if not already downloaded
nltk.download('punkt', quiet=True)

# ----------------------------
# Table Extraction Functions
# ----------------------------

def extract_tables_with_camelot(file_path: str) -> List[Dict]:
    """
    Extracts tables from a PDF file using Camelot.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        List[Dict]: A list of tables with page number, table number, and table data.
    """
    logger.info(f"Extracting tables from PDF using Camelot: {file_path}")
    tables_data = []
    try:
        # First attempt with 'lattice' flavor
        tables = camelot.read_pdf(file_path, pages='all', flavor='lattice')
        logger.info(f"Camelot found {tables.n} tables using 'lattice' flavor.")

        # If no tables found with 'lattice', try 'stream'
        if tables.n == 0:
            logger.info("No tables found with 'lattice' flavor. Trying 'stream' flavor...")
            tables = camelot.read_pdf(file_path, pages='all', flavor='stream')
            logger.info(f"Camelot found {tables.n} tables using 'stream' flavor.")

        for table_num, table in enumerate(tables, 1):
            df = table.df  # pandas DataFrame
            logger.debug(f"Extracted Table {table_num} on Page {table.page}:\n{df}")
            rows = []
            for _, row in df.iterrows():
                cells = row.tolist()
                sanitized_cells = [cell.strip() if cell.strip() else "" for cell in cells]
                rows.append({"cells": sanitized_cells})
            table_dict = {
                "page_number": table.page,  # Ensure this is always an integer
                "table_number": table_num,
                "rows": rows
            }
            tables_data.append(table_dict)
            # Detailed logging of each table's content
            logger.debug(f"Processed Table {table_num} on Page {table.page}:")
            for row in rows:
                logger.debug(f"Row: {row['cells']}")

    except Exception as e:
        logger.error(f"Failed to extract tables with Camelot: {e}")
        # Fallback to pdfplumber if Camelot fails
        logger.info("Falling back to pdfplumber for table extraction.")
        tables_data = extract_tables_with_pdfplumber(file_path)

    return tables_data

def extract_tables_with_pdfplumber(file_path: str) -> List[Dict]:
    """
    Extracts tables from a PDF file using pdfplumber as a fallback.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        List[Dict]: A list of tables with page number, table number, and table data.
    """
    logger.info(f"Extracting tables from PDF using pdfplumber: {file_path}")
    tables_data = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page_num, page in enumerate(pdf.pages, 1):
                tables = page.extract_tables()
                logger.info(f"pdfplumber - Page {page_num}: {len(tables)} tables found.")
                for i, table in enumerate(tables, 1):
                    # Ensure all cells are strings, replace None with empty string
                    sanitized_table = [
                        {"cells": [cell if cell is not None else "" for cell in row]}
                        for row in table
                    ]
                    table_dict = {
                        "page_number": page_num,  # Always an integer
                        "table_number": i,
                        "rows": sanitized_table
                    }
                    tables_data.append(table_dict)
    except Exception as e:
        logger.error(f"Failed to extract tables with pdfplumber: {e}")
    return tables_data

def extract_tables_with_tabula(file_path: str) -> List[Dict]:
    """
    Extracts tables from a PDF file using Tabula-py.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        List[Dict]: A list of tables with page number, table number, and table data.
    """
    logger.info(f"Extracting tables from PDF using Tabula-py: {file_path}")
    tables_data = []
    try:
        # Get total number of pages
        info = tabula.environment_info()
        total_pages = info['number_of_pages']

        for page_num in range(1, total_pages + 1):
            tables = tabula.read_pdf(file_path, pages=page_num, multiple_tables=True, silent=True)
            logger.info(f"Tabula-py - Page {page_num}: {len(tables)} tables found.")
            for table_num, df in enumerate(tables, 1):
                logger.debug(f"Extracted Table {table_num} on Page {page_num}:\n{df}")
                rows = []
                for _, row in df.iterrows():
                    cells = row.tolist()
                    sanitized_cells = [str(cell).strip() if pd.notna(cell) else "" for cell in cells]
                    rows.append({"cells": sanitized_cells})
                table_dict = {
                    "page_number": page_num,
                    "table_number": table_num,
                    "rows": rows
                }
                tables_data.append(table_dict)
    except Exception as e:
        logger.error(f"Failed to extract tables with Tabula-py: {e}")
    return tables_data

def extract_tables_with_python_docx(file_path: str) -> List[Dict]:
    """
    Extracts tables from a DOCX file using python-docx.

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        List[Dict]: A list of tables with table number and table data.
    """
    logger.info(f"Extracting tables from DOCX using python-docx: {file_path}")
    tables_data = []
    try:
        doc = docx.Document(file_path)
        for table_num, table in enumerate(doc.tables, 1):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip() if cell.text.strip() else "" for cell in row.cells]
                if any(cell for cell in cells):
                    rows.append({"cells": cells})
            table_dict = {
                "page_number": 0,  # DOCX doesn't have pages; set to 0 or another identifier
                "table_number": table_num,
                "rows": rows
            }
            tables_data.append(table_dict)
    except Exception as e:
        logger.error(f"Failed to extract tables from DOCX using python-docx: {e}")
    return tables_data

def extract_tables(file_path: str, file_type: str) -> Dict[str, List[Dict]]:
    """
    Extracts tables using multiple libraries/frameworks based on file type.

    Args:
        file_path (str): Path to the file.
        file_type (str): Type of the file ('pdf' or 'docx').

    Returns:
        Dict[str, List[Dict]]: Dictionary containing tables extracted by each method.
    """
    tables = {}
    if file_type == 'pdf':
        # Extract using Camelot
        camelot_tables = extract_tables_with_camelot(file_path)
        if camelot_tables:
            tables['camelot'] = camelot_tables

        # Extract using pdfplumber
        pdfplumber_tables = extract_tables_with_pdfplumber(file_path)
        if pdfplumber_tables:
            tables['pdfplumber'] = pdfplumber_tables

        # Extract using Tabula-py
        tabula_tables = extract_tables_with_tabula(file_path)
        if tabula_tables:
            tables['tabula-py'] = tabula_tables

    elif file_type == 'docx':
        # Extract using python-docx
        python_docx_tables = extract_tables_with_python_docx(file_path)
        if python_docx_tables:
            tables['python-docx'] = python_docx_tables

    # Add more file types and extraction methods if needed

    return tables

# ----------------------------
# Text Extraction Functions
# ----------------------------

def extract_text_from_pdf(file_path: str) -> Tuple[str, Dict[str, List[Dict]]]:
    """
    Extracts text and tables from a PDF file using PyMuPDF and multiple table extraction libraries.

    Args:
        file_path (str): Path to the PDF file.

    Returns:
        Tuple[str, Dict[str, List[Dict]]]: Extracted text and dictionary of tables extracted by different libraries.
    """
    logger.info(f"Extracting text from PDF: {file_path}")
    text = ""
    tables = {}
    try:
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc, 1):
                logger.info(f"Processing page {page_num}/{len(doc)}")
                page_text = page.get_text()
                if page_text.strip():
                    # Text extracted successfully
                    text += page_text + "\n"
                else:
                    # Fallback to OCR
                    logger.info(f"No text found on page {page_num}, using OCR.")
                    pix = page.get_pixmap()
                    image_bytes = pix.tobytes("png")
                    image = Image.open(io.BytesIO(image_bytes))
                    # Preprocess image
                    image = image.convert('L')  # Convert to grayscale
                    image = image.resize((image.width * 2, image.height * 2), Image.ANTIALIAS)
                    ocr_text = pytesseract.image_to_string(image, lang='eng', config='--psm 6')
                    text += ocr_text + "\n"

        # After extracting text, attempt to extract tables using multiple libraries
        tables = extract_tables(file_path, file_type='pdf')

    except Exception as e:
        logger.error(f"Failed to extract text from PDF: {e}")
        raise e
    return text, tables

def extract_text_from_docx(file_path: str) -> Tuple[str, Dict[str, List[Dict]]]:
    """
    Extracts text and tables from a DOCX file using python-docx.

    Args:
        file_path (str): Path to the DOCX file.

    Returns:
        Tuple[str, Dict[str, List[Dict]]]: Extracted text and dictionary of tables extracted by different methods.
    """
    logger.info(f"Extracting text and tables from DOCX: {file_path}")
    text = ""
    tables = {}
    try:
        doc = docx.Document(file_path)
        text = '\n'.join([para.text for para in doc.paragraphs])
        # Extract tables
        tables = extract_tables(file_path, file_type='docx')
    except Exception as e:
        logger.error(f"Failed to extract text from DOCX: {e}")
        raise e
    return text, tables

def extract_text_from_txt(file_path: str) -> Tuple[str, Dict[str, List[Dict]]]:
    """
    Extracts text from a TXT file.

    Args:
        file_path (str): Path to the TXT file.

    Returns:
        Tuple[str, Dict[str, List[Dict]]]: Extracted text and empty tables.
    """
    logger.info(f"Extracting text from TXT: {file_path}")
    text = ""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            text = f.read()
    except Exception as e:
        logger.error(f"Failed to extract text from TXT: {e}")
        raise e
    return text, {}

# ----------------------------
# Entity Extraction Function
# ----------------------------

def extract_entities_with_spacy(text: str) -> List[List[Dict]]:
    """
    Extracts entities from text using spaCy.

    Args:
        text (str): The input text.

    Returns:
        List[List[Dict]]: A list where each element corresponds to a chunk and contains a list of entities.
    """
    logger.info("Extracting entities from chunks using spaCy...")
    # Assuming chunk_text is already implemented to return chunks
    chunks = chunk_text(text, method='spacy')  # Adjust 'method' as needed
    entities_per_chunk = []
    for idx, chunk in enumerate(chunks, 1):
        logger.debug(f"Extracting entities from chunk {idx}/{len(chunks)} using spaCy.")
        try:
            doc = nlp(chunk)
            entities = [{"text": ent.text, "label": ent.label_} for ent in doc.ents]
            entities_per_chunk.append(entities)
        except Exception as e:
            logger.error(f"Failed to extract entities from chunk {idx}: {e}")
            entities_per_chunk.append([])
    logger.info("Entity extraction with spaCy completed.")
    return entities_per_chunk

# ----------------------------
# Metrics Computation Function
# ----------------------------

def compute_metrics(text: str, chunks: List[str]) -> Dict:
    """
    Computes various metrics based on the original text and its chunks.

    Args:
        text (str): The original extracted text.
        chunks (List[str]): The list of text chunks.

    Returns:
        Dict: A dictionary containing computed metrics.
    """
    logger.info("Computing metrics...")
    try:
        # Normalize newlines
        text = text.replace('\r\n', '\n').replace('\r', '\n')

        # Number of lines (non-empty lines)
        lines = text.split('\n')
        num_lines = len([line for line in lines if line.strip()])

        # Enhanced paragraph detection
        # Split paragraphs on newlines followed by an uppercase letter
        paragraph_pattern = re.compile(r'\n+(?=\s*[A-Z])')
        paragraphs = paragraph_pattern.split(text.strip())
        paragraphs = [para for para in paragraphs if para.strip()]
        num_paragraphs = len(paragraphs)

        # If no paragraphs detected, default to splitting on single newlines
        if num_paragraphs <= 1:
            paragraphs = [para for para in lines if para.strip()]
            num_paragraphs = len(paragraphs)

        # Number of words
        num_words = len(text.split())

        # Average words per paragraph
        avg_words_per_paragraph = num_words / num_paragraphs if num_paragraphs else 0

        # Average words per line
        avg_words_per_line = num_words / num_lines if num_lines else 0

        # Original content size in bytes
        original_content_size = len(text.encode('utf-8'))

        # Number of chunks
        num_chunks = len(chunks)

        metrics = {
            "num_lines": num_lines,
            "num_paragraphs": num_paragraphs,
            "num_words": num_words,
            "avg_words_per_paragraph": round(avg_words_per_paragraph, 2),
            "avg_words_per_line": round(avg_words_per_line, 2),
            "original_content_size": original_content_size,
            "num_chunks": num_chunks
        }

        logger.info(f"Metrics computed: {metrics}")
        return metrics
    except Exception as e:
        logger.error(f"Failed to compute metrics: {e}")
        raise e
